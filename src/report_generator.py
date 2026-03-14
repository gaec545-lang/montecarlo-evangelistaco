"""
Report Generator - FASE 5
Evangelista & Co. | Sentinel Decision Intelligence V2

Genera reportes corporativos en PDF y DOCX con:
  - Portada con logo Evangelista & Co.
  - Resumen ejecutivo
  - Resultados de los 3 Escudos
  - Recomendaciones y plan de rescate
"""

import os
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd

logger = logging.getLogger(__name__)

# Ruta al logo (relativa a la raíz del proyecto)
LOGO_PATH = Path(__file__).resolve().parent.parent / "assets" / "logoEvangelistaCo.png"

# Paleta corporativa
COLOR_DORADO  = (212, 175, 55)    # #D4AF37
COLOR_MARINO  = (26,  26,  46)    # #1A1A2E
COLOR_BLANCO  = (255, 255, 255)
COLOR_GRIS    = (245, 245, 245)
COLOR_ROJO    = (198, 40,  40)
COLOR_VERDE   = (46,  125, 50)
COLOR_AMARILLO = (245, 127, 23)


# ── PDF ────────────────────────────────────────────────────────────────────────

def _color_rgb(r, g, b):
    """Convierte RGB 0-255 a ReportLab 0-1."""
    from reportlab.lib.colors import Color
    return Color(r / 255, g / 255, b / 255)


def _build_pdf(client_name: str, results: dict, output_path: str) -> str:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, Image, PageBreak, KeepTogether,
    )
    from reportlab.lib import colors

    DORADO   = _color_rgb(*COLOR_DORADO)
    MARINO   = _color_rgb(*COLOR_MARINO)
    ROJO_RL  = _color_rgb(*COLOR_ROJO)
    VERDE_RL = _color_rgb(*COLOR_VERDE)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.75 * inch, leftMargin=0.75 * inch,
        topMargin=0.75 * inch,   bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    style_titulo    = ParagraphStyle("titulo",    parent=styles["Title"],
                                      textColor=MARINO, fontSize=22, spaceAfter=6, alignment=TA_CENTER)
    style_subtitulo = ParagraphStyle("subtitulo", parent=styles["Normal"],
                                      textColor=DORADO, fontSize=13, spaceAfter=4, alignment=TA_CENTER)
    style_h1        = ParagraphStyle("h1",        parent=styles["Heading1"],
                                      textColor=MARINO, fontSize=14, spaceBefore=14, spaceAfter=6)
    style_h2        = ParagraphStyle("h2",        parent=styles["Heading2"],
                                      textColor=DORADO, fontSize=11, spaceBefore=8,  spaceAfter=4)
    style_body      = ParagraphStyle("body",      parent=styles["Normal"],
                                      fontSize=10, spaceAfter=6, leading=14)
    style_pie       = ParagraphStyle("pie",       parent=styles["Normal"],
                                      fontSize=8,  textColor=colors.grey, alignment=TA_CENTER)
    style_alerta    = ParagraphStyle("alerta",    parent=styles["Normal"],
                                      fontSize=10, textColor=ROJO_RL, spaceAfter=4)
    style_ok        = ParagraphStyle("ok",        parent=styles["Normal"],
                                      fontSize=10, textColor=VERDE_RL, spaceAfter=4)

    story = []
    fecha_hoy = datetime.now().strftime("%d de %B de %Y")

    # ── PORTADA ────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.8 * inch))

    if LOGO_PATH.exists():
        story.append(Image(str(LOGO_PATH), width=2.2 * inch, height=1.1 * inch))
        story.append(Spacer(1, 0.3 * inch))

    story.append(Paragraph("DICTAMEN DE INTELIGENCIA DE DECISIONES", style_titulo))
    story.append(Paragraph(client_name, style_subtitulo))
    story.append(Spacer(1, 0.15 * inch))
    story.append(HRFlowable(width="100%", thickness=2, color=DORADO, spaceAfter=8))
    story.append(Paragraph(f"Fecha de Emisión: {fecha_hoy}", style_pie))
    story.append(Paragraph("Confidencial — Exclusivo para uso interno", style_pie))
    story.append(Spacer(1, 0.3 * inch))

    # Semáforo ejecutivo en portada
    stress = results.get("stress_results", {})
    prob   = stress.get("probabilidad_crisis", 0)
    mes_c  = stress.get("mes_critico")

    if prob > 0.30:
        semaforo_txt = f"🔴  ALERTA: Probabilidad de crisis {prob:.0%} — Mes crítico: {mes_c}"
        story.append(Paragraph(semaforo_txt, style_alerta))
    elif prob > 0.15:
        story.append(Paragraph(f"🟡  PRECAUCIÓN: Probabilidad de crisis {prob:.0%}", style_body))
    else:
        story.append(Paragraph("🟢  SALUDABLE: Sin crisis proyectada en el horizonte", style_ok))

    story.append(PageBreak())

    # ── RESUMEN EJECUTIVO ──────────────────────────────────────────────────────
    story.append(Paragraph("1. RESUMEN EJECUTIVO", style_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=DORADO, spaceAfter=6))

    narrative = results.get("business_narrative", {})
    if isinstance(narrative, dict):
        exec_summary = narrative.get("executive_summary", "")
        if exec_summary:
            story.append(Paragraph(exec_summary, style_body))

    # KPIs de Monte Carlo
    stats = results.get("statistics", {})
    if stats:
        story.append(Spacer(1, 0.15 * inch))
        story.append(Paragraph("Indicadores Clave de Riesgo (Monte Carlo 10,000 sim.)", style_h2))
        kpi_data = [
            ["Indicador", "Valor"],
            ["Resultado Esperado (P50)",   f"${stats.get('p50', 0):,.0f} MXN"],
            ["Escenario Optimista (P90)",  f"${stats.get('p90', 0):,.0f} MXN"],
            ["Escenario Pesimista (P10)",  f"${stats.get('p10', 0):,.0f} MXN"],
            ["Probabilidad de Pérdida",    f"{stats.get('prob_loss', 0):.1%}"],
            ["VaR 95%",                    f"${abs(stats.get('var_95', 0)):,.0f} MXN"],
        ]
        kpi_table = Table(kpi_data, colWidths=[3.5 * inch, 3 * inch])
        kpi_table.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, 0), MARINO),
            ("TEXTCOLOR",    (0, 0), (-1, 0), _color_rgb(*COLOR_BLANCO)),
            ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, -1), 10),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_color_rgb(*COLOR_GRIS), _color_rgb(*COLOR_BLANCO)]),
            ("GRID",         (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("PADDING",      (0, 0), (-1, -1), 6),
        ]))
        story.append(kpi_table)

    story.append(PageBreak())

    # ── ESCUDO 1: PROYECCIONES ─────────────────────────────────────────────────
    story.append(Paragraph("2. ESCUDO 1 — RADAR: Proyecciones Base", style_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=DORADO, spaceAfter=6))

    forecasting = results.get("forecasting_results", {})
    if forecasting and "error" not in forecasting:
        df_flujo = forecasting.get("flujo_libre_12m")
        if df_flujo is not None and not df_flujo.empty:
            story.append(Paragraph("Proyección de Flujo Libre a 12 Meses", style_h2))
            flujo_data = [["Mes", "Ingresos Proy.", "Costos Proy.", "Flujo Libre", "Flujo Acum."]]
            for _, row in df_flujo.iterrows():
                flujo_data.append([
                    str(row["mes"]) if "mes" in row.index else row["fecha"].strftime("%b %Y"),
                    f"${row['ingresos']:,.0f}",
                    f"${row['costos']:,.0f}",
                    f"${row['flujo_libre']:,.0f}",
                    f"${row['flujo_acumulado']:,.0f}",
                ])
            t = Table(flujo_data, colWidths=[1 * inch, 1.5 * inch, 1.5 * inch, 1.5 * inch, 1.5 * inch])
            t.setStyle(TableStyle([
                ("BACKGROUND",   (0, 0), (-1, 0), MARINO),
                ("TEXTCOLOR",    (0, 0), (-1, 0), _color_rgb(*COLOR_BLANCO)),
                ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",     (0, 0), (-1, -1), 8),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_color_rgb(*COLOR_GRIS), _color_rgb(*COLOR_BLANCO)]),
                ("GRID",         (0, 0), (-1, -1), 0.3, colors.lightgrey),
                ("PADDING",      (0, 0), (-1, -1), 4),
                ("ALIGN",        (1, 0), (-1, -1), "RIGHT"),
            ]))
            story.append(t)

        estacional = forecasting.get("estacionalidad_detectada", {})
        if estacional.get("detectada"):
            story.append(Spacer(1, 0.1 * inch))
            story.append(Paragraph(
                f"Estacionalidad detectada: Mes pico → {estacional.get('mes_pico')} | "
                f"Mes valle → {estacional.get('mes_valle')}",
                style_body
            ))
    else:
        story.append(Paragraph("Sin datos de proyección disponibles.", style_body))

    story.append(PageBreak())

    # ── ESCUDO 2: ANÁLISIS DE RIESGO ──────────────────────────────────────────
    story.append(Paragraph("3. ESCUDO 2 — TRITURADORA: Análisis de Riesgo", style_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=DORADO, spaceAfter=6))

    if stress and "error" not in stress:
        riesgo_data = [
            ["Métrica", "Valor"],
            ["Probabilidad de Crisis de Liquidez",  f"{stress.get('probabilidad_crisis', 0):.1%}"],
            ["Mes Crítico",                          str(stress.get('mes_critico', 'N/A'))],
            ["Evento Detonante",                     stress.get('evento_detonante', 'N/A')],
            ["Prob. Default de Clientes (bayesiano)", f"{stress.get('default_probability', {}).get('prob_default_media', 0):.1%}"],
        ]
        perc = stress.get("percentiles_caja", {})
        if perc:
            riesgo_data += [
                ["Caja Proyectada P10 (pesimista)", f"${perc.get('p10', 0):,.0f} MXN"],
                ["Caja Proyectada P50 (base)",      f"${perc.get('p50', 0):,.0f} MXN"],
                ["Caja Proyectada P90 (optimista)", f"${perc.get('p90', 0):,.0f} MXN"],
            ]
        t2 = Table(riesgo_data, colWidths=[3.5 * inch, 3 * inch])
        t2.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, 0), MARINO),
            ("TEXTCOLOR",    (0, 0), (-1, 0), _color_rgb(*COLOR_BLANCO)),
            ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, -1), 10),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_color_rgb(*COLOR_GRIS), _color_rgb(*COLOR_BLANCO)]),
            ("GRID",         (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("PADDING",      (0, 0), (-1, -1), 6),
        ]))
        story.append(t2)

    story.append(PageBreak())

    # ── ESCUDO 3: PLAN DE RESCATE ──────────────────────────────────────────────
    story.append(Paragraph("4. ESCUDO 3 — BISTURÍ: Plan de Rescate Óptimo", style_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=DORADO, spaceAfter=6))

    opt = results.get("optimization_results", {})
    if opt and not opt.get("error"):
        if opt.get("crisis_detectada"):
            story.append(Paragraph(
                f"Capital Total a Liberar: ${opt.get('capital_total_liberado', 0):,.0f} MXN  |  "
                f"ROI Estimado del Plan: {opt.get('roi_estimado', 0)}x",
                style_h2
            ))
            for i, est in enumerate(opt.get("estrategias", []), 1):
                story.append(Spacer(1, 0.1 * inch))
                story.append(Paragraph(f"Estrategia #{i}: {est.get('titulo', '')}", style_h2))
                est_data = [
                    ["Acción", est.get("accion", "")],
                    ["Capital Liberado", f"${est.get('capital_liberado', 0):,.0f} MXN"],
                    ["Deadline", est.get("deadline", "")],
                ]
                te = Table(est_data, colWidths=[1.8 * inch, 4.7 * inch])
                te.setStyle(TableStyle([
                    ("FONTNAME",  (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE",  (0, 0), (-1, -1), 10),
                    ("GRID",      (0, 0), (-1, -1), 0.3, colors.lightgrey),
                    ("PADDING",   (0, 0), (-1, -1), 5),
                    ("VALIGN",    (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(te)
                story.append(Paragraph(est.get("descripcion", ""), style_body))
        else:
            story.append(Paragraph("✅ " + opt.get("mensaje", "No se detectó crisis."), style_ok))

    # ── RECOMENDACIONES DE DECISION INTELLIGENCE ──────────────────────────────
    recs = results.get("recommendations", [])
    if recs:
        story.append(PageBreak())
        story.append(Paragraph("5. INTELIGENCIA DE DECISIONES", style_h1))
        story.append(HRFlowable(width="100%", thickness=1, color=DORADO, spaceAfter=6))
        for i, rec in enumerate(recs[:3], 1):
            story.append(Paragraph(
                f"#{i} {rec.get('title', '')} — Prioridad {rec.get('priority', '')}",
                style_h2
            ))
            story.append(Paragraph(rec.get("description", ""), style_body))

    # ── PIE DE PÁGINA FINAL ────────────────────────────────────────────────────
    story.append(Spacer(1, 0.5 * inch))
    story.append(HRFlowable(width="100%", thickness=1, color=DORADO, spaceAfter=4))
    story.append(Paragraph(
        f"© {datetime.now().year} Evangelista & Co. — Sentinel Decision Intelligence V2 | "
        f"Documento generado el {fecha_hoy}",
        style_pie
    ))

    doc.build(story)
    logger.info(f"PDF generado: {output_path}")
    return output_path


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _build_docx(client_name: str, results: dict, output_path: str) -> str:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    DORADO_DOCX = RGBColor(*COLOR_DORADO)
    MARINO_DOCX = RGBColor(*COLOR_MARINO)

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    def add_heading(text, level=1, color=None):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if color:
            for run in h.runs:
                run.font.color.rgb = color
        return h

    def add_para(text, bold=False, italic=False, color=None, size=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
        return p

    def add_table_2col(data: list, header_color=None):
        t = doc.add_table(rows=len(data), cols=2)
        t.style = "Table Grid"
        for i, (k, v) in enumerate(data):
            cells = t.rows[i].cells
            cells[0].text = str(k)
            cells[1].text = str(v)
            if i == 0 and header_color:
                for cell in cells:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "1A1A2E")
                    cell._tc.get_or_add_tcPr().append(shading)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
        return t

    # ── Portada
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(2.2))

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("DICTAMEN DE INTELIGENCIA DE DECISIONES")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = MARINO_DOCX

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(client_name)
    run2.font.size = Pt(14)
    run2.font.color.rgb = DORADO_DOCX

    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_p.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')} | Confidencial")

    doc.add_page_break()

    # ── Resumen Ejecutivo
    add_heading("1. RESUMEN EJECUTIVO", 1, MARINO_DOCX)
    narrative = results.get("business_narrative", {})
    if isinstance(narrative, dict) and narrative.get("executive_summary"):
        add_para(narrative["executive_summary"])

    stats = results.get("statistics", {})
    if stats:
        add_heading("Indicadores Monte Carlo (10,000 simulaciones)", 2, DORADO_DOCX)
        add_table_2col([
            ["Indicador", "Valor"],
            ["P50 (Resultado Esperado)", f"${stats.get('p50', 0):,.0f} MXN"],
            ["P90 (Optimista)",          f"${stats.get('p90', 0):,.0f} MXN"],
            ["P10 (Pesimista)",          f"${stats.get('p10', 0):,.0f} MXN"],
            ["Prob. de Pérdida",         f"{stats.get('prob_loss', 0):.1%}"],
        ], header_color=True)

    doc.add_page_break()

    # ── Escudos
    add_heading("2. ESCUDO 1 — Proyecciones (Radar)", 1, MARINO_DOCX)
    forecasting = results.get("forecasting_results", {})
    if forecasting and "error" not in forecasting:
        df_flujo = forecasting.get("flujo_libre_12m")
        if df_flujo is not None and not df_flujo.empty:
            rows = [["Fecha", "Ingresos", "Costos", "Flujo Libre", "Acumulado"]]
            for _, row in df_flujo.iterrows():
                fecha_lbl = str(row.get("mes", row.get("fecha", "")))
                rows.append([
                    fecha_lbl,
                    f"${row['ingresos']:,.0f}",
                    f"${row['costos']:,.0f}",
                    f"${row['flujo_libre']:,.0f}",
                    f"${row['flujo_acumulado']:,.0f}",
                ])
            t = doc.add_table(rows=len(rows), cols=5)
            t.style = "Table Grid"
            for i, row_data in enumerate(rows):
                for j, val in enumerate(row_data):
                    t.rows[i].cells[j].text = val
    else:
        add_para("Sin datos de proyección disponibles.")

    add_heading("3. ESCUDO 2 — Análisis de Riesgo (Trituradora)", 1, MARINO_DOCX)
    stress = results.get("stress_results", {})
    if stress and "error" not in stress:
        add_table_2col([
            ["Métrica", "Valor"],
            ["Prob. Crisis",        f"{stress.get('probabilidad_crisis', 0):.1%}"],
            ["Mes Crítico",         str(stress.get('mes_critico', 'N/A'))],
            ["Evento Detonante",    stress.get('evento_detonante', 'N/A')],
            ["Caja P50",            f"${stress.get('percentiles_caja', {}).get('p50', 0):,.0f} MXN"],
        ], header_color=True)

    doc.add_page_break()

    add_heading("4. ESCUDO 3 — Plan de Rescate (Bisturí)", 1, MARINO_DOCX)
    opt = results.get("optimization_results", {})
    if opt and not opt.get("error"):
        if opt.get("crisis_detectada"):
            add_para(
                f"Capital Total: ${opt.get('capital_total_liberado', 0):,.0f} MXN  |  "
                f"ROI: {opt.get('roi_estimado', 0)}x",
                bold=True, color=DORADO_DOCX
            )
            for i, est in enumerate(opt.get("estrategias", []), 1):
                add_heading(f"Estrategia #{i}: {est.get('titulo', '')}", 2, DORADO_DOCX)
                add_table_2col([
                    ["Acción",           est.get("accion", "")],
                    ["Capital Liberado", f"${est.get('capital_liberado', 0):,.0f} MXN"],
                    ["Deadline",         est.get("deadline", "")],
                ])
                add_para(est.get("descripcion", ""))
        else:
            add_para("✅ " + opt.get("mensaje", "No se detectó crisis."))

    # Pie
    doc.add_paragraph()
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(
        f"© {datetime.now().year} Evangelista & Co. — Sentinel Decision Intelligence V2"
    )
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
    logger.info(f"DOCX generado: {output_path}")
    return output_path


# ── API pública ────────────────────────────────────────────────────────────────

class ReportGenerator:
    """
    Genera reportes en PDF y DOCX con branding corporativo Evangelista & Co.

    Uso:
        gen = ReportGenerator(client_name="Cibrián Arquitectos", results=pipeline_results)
        pdf_path  = gen.generate_pdf("/tmp/reporte.pdf")
        docx_path = gen.generate_docx("/tmp/reporte.docx")
    """

    def __init__(self, client_name: str, results: dict):
        self.client_name = client_name
        self.results     = results

    def generate_pdf(self, output_path: str) -> str:
        return _build_pdf(self.client_name, self.results, output_path)

    def generate_docx(self, output_path: str) -> str:
        return _build_docx(self.client_name, self.results, output_path)

    def generate_both(self, output_dir: str, prefix: str = "Sentinel_Reporte") -> dict:
        """Genera PDF y DOCX en el directorio indicado."""
        os.makedirs(output_dir, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M")
        pdf_path  = os.path.join(output_dir, f"{prefix}_{stamp}.pdf")
        docx_path = os.path.join(output_dir, f"{prefix}_{stamp}.docx")
        return {
            "pdf":  self.generate_pdf(pdf_path),
            "docx": self.generate_docx(docx_path),
        }
